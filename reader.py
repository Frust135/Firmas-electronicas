from shutil import copyfile #Función para copiar archivos
from docx import Document #Pip install python-docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import filedialog
import docx

#-----------------------------------------
#         Lector de USB
#-----------------------------------------

def usbReader():
    clave_privada = None
    try:
        with open('E:/token.txt', 'r') as file:
            clave_privada = file.read()
            return clave_privada
    except:
        return None
        
#-----------------------------------------
#         Firma el documento
#-----------------------------------------

def estamparDocumento(directorio_documento, nombre_archivo, estampado):
    try:
        path = r'./files/'+nombre_archivo+'.docx'
        documento = Document(directorio_documento)
        firma = documento.add_paragraph() #Crea un párrafo
        firma.alignment=  WD_PARAGRAPH_ALIGNMENT.RIGHT #Y lo alinea a la derecha
        ejecutar = firma.add_run("") 
        ejecutar.add_picture(estampado, width=Inches(1.75)) #Y se añade la imagen al párrafo
        documento.save(path)
    except:
        print("Error al firmar el documento.")

#-----------------------------------------
#         Obtener texto de todo el documento
#-----------------------------------------
def getText(path):
    doc = docx.Document(path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText

#-----------------------------------------
#         Obtener un Documento en TKINTER
#-----------------------------------------

def openFile(ventana):
    ventana.filename = filedialog.askopenfilename(initialdir='/', 
                                title="Selecciona el Archivo",
                                filetypes=(("Archivos Docx","*.docx"),
                                        ("Todos los archivos","*.*")))
    return ventana.filename

